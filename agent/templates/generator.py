from __future__ import annotations

from pathlib import Path
from typing import Dict, Any, List, Optional

from agent.parser.router import FileRouter
from agent.parser.schema import ParsedDocument

class TemplateGenerator:
    def __init__(self, templates_dir: str = "data/templates"):
        # 确保模板目录路径是相对于项目根目录的
        self.templates_dir = Path(__file__).resolve().parents[2] / templates_dir
        self.router = FileRouter()
    
    def generate_from_template(
        self,
        template_name: str,
        data: Dict[str, Any],
        output_path: Optional[str] = None,
        output_dir: Optional[str] = None
    ) -> Dict[str, Any]:
        template_path = self.templates_dir / template_name
        if not template_path.exists():
            return {
                "success": False,
                "error": f"Template not found: {template_path}"
            }
        
        if output_path is None:
            from datetime import datetime
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            extension = template_path.suffix
            
            # 如果指定了输出目录
            if output_dir:
                output_path = f"{output_dir}/{template_name.replace(extension, '')}_{timestamp}{extension}"
            else:
                output_path = f"docs/parsed/{template_name.replace(extension, '')}_{timestamp}{extension}"
        
        # 对于Word文档，直接基于原始模板进行替换
        if template_path.suffix == ".docx":
            return self._generate_word_from_template(str(template_path), data, output_path)
        # 对于Excel文档，使用模板替换方式
        elif template_path.suffix in [".xlsx", ".xls"]:
            return self._generate_excel_from_template(str(template_path), data, output_path)
        else:
            # 对于其他格式，使用解析后重新生成的方式
            parsed_template = self.router.parse_file(str(template_path))
            filled_doc = self._fill_data(parsed_template, data)
            return self._generate_document(filled_doc, output_path)
    
    def _generate_excel_from_template(self, template_path: str, data: Dict[str, Any], output_path: str) -> Dict[str, Any]:
        """
        基于 Excel 模板生成文档
        支持多条数据的填充，使用 {{item.field}} 占位符和 {% for item in data_list %} 循环标记
        """
        try:
            from openpyxl import load_workbook
            import re
        except ImportError:
            return {
                "success": False,
                "error": "openpyxl library not installed"
            }
        
        try:
            # 加载模板
            wb = load_workbook(template_path)
            ws = wb.active
            
            # 打印模板信息
            print(f"模板路径: {template_path}")
            print(f"工作表名称: {ws.title}")
            print(f"最大行数: {ws.max_row}")
            print(f"最大列数: {ws.max_column}")
            
            # 查找循环标记
            data_list_name = "payment_data_list"  # 默认数据列表名称
            template_rows = []
            template_area_start = None
            template_area_end = None
            
            # 查找模板行（包含 {{item.}} 占位符的行）
            for row in range(1, ws.max_row + 1):
                has_placeholder = False
                row_data = []
                for col in range(1, ws.max_column + 1):
                    cell = ws.cell(row=row, column=col)
                    if cell.value:
                        cell_value = str(cell.value)
                        if '{{item.' in cell_value:
                            has_placeholder = True
                            print(f"找到占位符: {cell_value} 在行 {row}, 列 {col}")
                    row_data.append({
                        "value": cell.value
                    })
                if has_placeholder:
                    if template_area_start is None:
                        template_area_start = row
                    template_area_end = row
                    template_rows.append(row_data)
            
            # 查找循环结束标记
            end_row = None
            if template_area_end:
                for row in range(template_area_end, ws.max_row + 1):
                    for col in range(1, ws.max_column + 1):
                        cell = ws.cell(row=row, column=col)
                        if cell.value and '{% endfor %}' in str(cell.value):
                            end_row = row
                            print(f"找到结束标记在行 {row}")
                            break
                    if end_row:
                        break
            
            # 如果找到模板行和结束标记
            if template_rows and end_row:
                print(f"模板区域: 行 {template_area_start} 到 {template_area_end}")
                print(f"结束标记在: 行 {end_row}")
                print(f"模板行数: {len(template_rows)}")
                
                # 清空模板区域到结束标记之间的内容
                for row in range(template_area_start, end_row + 1):
                    for col in range(1, ws.max_column + 1):
                        ws.cell(row=row, column=col).value = None
                
                # 填充数据
                data_list = data.get(data_list_name, [])
                print(f"数据列表长度: {len(data_list)}")
                
                current_row = template_area_start
                
                for item_idx, item in enumerate(data_list):
                    print(f"填充第 {item_idx + 1} 条数据: {item}")
                    for template_row in template_rows:
                        for col_idx, cell_template in enumerate(template_row):
                            col = col_idx + 1
                            cell = ws.cell(row=current_row, column=col)
                            
                            # 替换占位符
                            if cell_template['value']:
                                cell_value = str(cell_template['value'])
                                replaced_value = self._replace_placeholders(cell_value, item)
                                cell.value = replaced_value
                                print(f"替换占位符: {cell_value} -> {replaced_value} 在行 {current_row}, 列 {col}")
                        
                        current_row += 1
            else:
                # 没有找到循环标记，使用普通替换
                print("没有找到循环标记，使用普通替换")
                for row in range(1, ws.max_row + 1):
                    for col in range(1, ws.max_column + 1):
                        cell = ws.cell(row=row, column=col)
                        if cell.value:
                            cell.value = self._replace_placeholders(str(cell.value), data)
            
            # 保存
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            wb.save(str(output_path))
            print(f"保存文件到: {output_path}")
            
            return {
                "success": True,
                "output_path": str(output_path),
                "sheet_name": ws.title
            }
            
        except Exception as e:
            print(f"错误: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _fill_data(self, doc: ParsedDocument, data: Dict[str, Any]) -> ParsedDocument:
        for section in doc.sections:
            section.content = self._replace_placeholders(section.content, data)
        
        for table in doc.tables:
            filled_rows = []
            for row in table.rows:
                filled_row = []
                for cell in row:
                    if isinstance(cell, str):
                        filled_row.append(self._replace_placeholders(cell, data))
                    else:
                        filled_row.append(cell)
                filled_rows.append(filled_row)
            table.rows = filled_rows
        
        return doc
    
    def _replace_placeholders(self, text: str, data: Dict[str, Any]) -> str:
        import re
        
        # 替换 {{item.field}} 格式
        def replace_item_placeholder(match):
            field = match.group(1)
            return str(data.get(field, match.group(0)))
        
        text = re.sub(r'\{\{item\.([^}]+)\}\}', replace_item_placeholder, text)
        
        # 替换 {{field}} 格式
        def replace_placeholder(match):
            field = match.group(1)
            return str(data.get(field, match.group(0)))
        
        text = re.sub(r'\{\{([\w\u4e00-\u9fa5]+)\}\}', replace_placeholder, text)
        
        # 替换 [field] 格式
        def replace_bracket_placeholder(match):
            field = match.group(1)
            return str(data.get(field, match.group(0)))
        
        text = re.sub(r'\[([\w\u4e00-\u9fa5]+)\]', replace_bracket_placeholder, text)
        
        return text
    
    def _generate_document(self, doc: ParsedDocument, output_path: str) -> Dict[str, Any]:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        if doc.content_type == "word":
            return self._generate_word(doc, output_path)
        elif doc.content_type == "excel":
            return self._generate_excel(doc, output_path)
        
        return {"success": False, "error": "Unsupported content type"}
    
    def _generate_word_from_template(self, template_path: str, data: Dict[str, Any], output_path: str) -> Dict[str, Any]:
        """
        直接基于原始Word模板进行替换，保持模板格式完整
        """
        try:
            from docx import Document
        except ImportError:
            return {
                "success": False,
                "error": "python-docx library not installed"
            }
        
        # 打开原始模板
        doc = Document(template_path)
        
        # 替换段落中的占位符（保持格式）
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.text:
                    run.text = self._replace_placeholders(run.text, data)
        
        # 替换表格中的占位符（保持格式）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text:
                                run.text = self._replace_placeholders(run.text, data)
        
        # 保存为新文件
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        return {
            "success": True,
            "output_path": str(output_path),
            "template_used": template_path
        }
    
    def _generate_word(self, doc: ParsedDocument, output_path: Path) -> Dict[str, Any]:
        """
        基于解析结果生成Word文档（备用方法）
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            return {
                "success": False,
                "error": "python-docx library not installed"
            }
        
        new_doc = Document()
        
        for section in doc.sections:
            if section.heading:
                heading_level = min(section.level, 9)
                new_doc.add_heading(section.heading, level=heading_level)
            
            if section.content:
                new_doc.add_paragraph(section.content)
        
        for table in doc.tables:
            if not table.headers:
                continue
            
            word_table = new_doc.add_table(
                rows=len(table.rows) + 1,
                cols=len(table.headers)
            )
            
            for col, header in enumerate(table.headers):
                word_table.rows[0].cells[col].text = header
            
            for row_idx, row in enumerate(table.rows, start=1):
                for col_idx, cell in enumerate(row):
                    if col_idx < len(table.headers):
                        word_table.rows[row_idx].cells[col_idx].text = str(cell) if cell else ""
        
        new_doc.save(str(output_path))
        
        return {
            "success": True,
            "output_path": str(output_path),
            "sections_count": len(doc.sections),
            "tables_count": len(doc.tables)
        }
    
    def _generate_excel(self, doc: ParsedDocument, output_path: Path) -> Dict[str, Any]:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
        except ImportError:
            return {
                "success": False,
                "error": "openpyxl library not installed"
            }
        
        wb = Workbook()
        ws = wb.active
        
        for table_idx, table in enumerate(doc.tables):
            if table_idx > 0:
                ws = wb.create_sheet(title=f"Sheet{table_idx + 1}")
            
            if table.headers:
                for col, header in enumerate(table.headers, start=1):
                    cell = ws.cell(row=1, column=col, value=header)
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                
                for row_idx, row in enumerate(table.rows, start=2):
                    for col_idx, cell in enumerate(row, start=1):
                        if cell is not None:
                            ws.cell(row=row_idx, column=col_idx, value=cell)
        
        wb.save(str(output_path))
        
        return {
            "success": True,
            "output_path": str(output_path),
            "tables_count": len(doc.tables)
        }