"""
agent/parser/parsers/docx_parser.py

Word 解析器。
- 使用 python-docx 提取段落（保留标题样式层级）、表格、图片
- 使用 mammoth 辅助生成 Markdown（可选）
- 每个段落记录 loc（paragraph 序号）
- 表格输出为 TableBlock，含坐标映射
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from agent.parser.base import BaseParser
from agent.parser.schema import (
    Error, ImageBlock, Loc, LocType, ParsedDocument, Section,
    TableBlock, TableMeta, Warning,
)
from agent.parser.utils.hash_utils import file_sha1

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    supported_extensions = [".docx"]

    def __init__(self, kb_name: str = ""):
        self.kb_name = kb_name

    def parse(self, file_path: str, **kwargs) -> ParsedDocument:
        from docx import Document as DocxDocument
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        doc_id = file_sha1(file_path, prefix=self.kb_name)
        fpath = Path(file_path)
        warnings: list[Warning] = []
        errors: list[Error] = []

        try:
            doc = DocxDocument(file_path)
        except Exception as exc:
            return ParsedDocument(
                doc_id=doc_id,
                source=self._source_dict(fpath),
                content_type="word",
                errors=[Error(
                    error_code="DOCX_OPEN_FAILED",
                    message=str(exc),
                    loc=Loc(type=LocType.PARAGRAPH.value, value=0),
                )],
            )

        # ---- 提取段落 → sections ----
        sections = self._extract_sections(doc, warnings)

        # ---- 提取表格 ----
        tables = self._extract_tables(doc, warnings)

        # ---- 提取图片 ----
        images = self._extract_images(doc, fpath, warnings)

        # ---- 标题 ----
        title = self._detect_title(doc)

        # ---- 核心属性 ----
        meta = {}
        try:
            cp = doc.core_properties
            meta["author"] = cp.author or ""
            meta["created"] = str(cp.created) if cp.created else ""
            meta["modified"] = str(cp.modified) if cp.modified else ""
            meta["subject"] = cp.subject or ""
        except Exception:
            pass
        meta["total_paragraphs"] = len(doc.paragraphs)
        meta["total_tables"] = len(doc.tables)

        return ParsedDocument(
            doc_id=doc_id,
            source=self._source_dict(fpath),
            title=title,
            content_type="word",
            sections=sections,
            tables=tables,
            images=images,
            warnings=warnings,
            errors=errors,
            metadata=meta,
        )

    # ------------------------------------------------------------------
    # 段落 → Sections
    # ------------------------------------------------------------------
    def _extract_sections(
        self, doc, warnings: list[Warning]
    ) -> list[Section]:
        sections: list[Section] = []
        current_heading = ""
        current_level = 1
        current_lines: list[str] = []
        current_para_idx = 0

        for para_idx, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            heading_level = self._style_to_level(style_name)

            if heading_level is not None and text:
                # 保存上一段
                if current_heading or current_lines:
                    sections.append(Section(
                        heading=current_heading,
                        level=current_level,
                        content="\n".join(current_lines).strip(),
                        loc=Loc(type=LocType.PARAGRAPH.value, value=current_para_idx),
                    ))
                current_heading = text
                current_level = heading_level
                current_lines = []
                current_para_idx = para_idx
            else:
                if text:
                    current_lines.append(text)
                elif current_lines and current_lines[-1] != "":
                    current_lines.append("")  # 保留段落间空行

        # 最后一段
        if current_heading or current_lines:
            sections.append(Section(
                heading=current_heading,
                level=current_level,
                content="\n".join(current_lines).strip(),
                loc=Loc(type=LocType.PARAGRAPH.value, value=current_para_idx),
            ))

        # 如果未检测到任何标题样式，发出 warning
        if all(s.heading == "" for s in sections):
            warnings.append(Warning(
                code="NO_HEADING_STYLES",
                message="No heading styles detected in document; structure may be flat.",
                loc=Loc(type=LocType.PARAGRAPH.value, value=0),
            ))

        return sections

    @staticmethod
    def _style_to_level(style_name: str) -> int | None:
        """Word 样式名 → 标题层级"""
        if not style_name:
            return None
        name_lower = style_name.lower()
        # "Heading 1" / "标题 1" 等
        m = re.match(r'heading\s*(\d)', name_lower)
        if m:
            return int(m.group(1))
        # 中文样式
        if "标题" in style_name:
            m2 = re.search(r'(\d)', style_name)
            if m2:
                return int(m2.group(1))
            if "标题" == style_name.strip():
                return 1
        if name_lower in ("title",):
            return 1
        if name_lower in ("subtitle",):
            return 2
        return None

    # ------------------------------------------------------------------
    # 表格
    # ------------------------------------------------------------------
    def _extract_tables(
        self, doc, warnings: list[Warning]
    ) -> list[TableBlock]:
        tables: list[TableBlock] = []
        for ti, table in enumerate(doc.tables):
            try:
                table_id = f"t{ti + 1}"
                all_rows = []
                for row in table.rows:
                    all_rows.append([cell.text.strip() for cell in row.cells])

                if not all_rows:
                    warnings.append(Warning(
                        code="EMPTY_TABLE",
                        message=f"Table {ti+1} is empty.",
                        loc=Loc(type=LocType.PARAGRAPH.value, value=0,
                                extra={"table_index": ti}),
                    ))
                    continue

                headers = all_rows[0]
                data_rows = all_rows[1:]

                # 检测合并单元格（python-docx 中合并单元格会导致重复文本）
                merged = self._detect_merged_cells(table)
                if merged:
                    warnings.append(Warning(
                        code="MERGED_CELLS_DETECTED",
                        message=f"Table {ti+1} has {len(merged)} merged cell region(s). Values may be duplicated.",
                        loc=Loc(type=LocType.PARAGRAPH.value, value=0,
                                extra={"table_index": ti, "merged_regions": merged}),
                    ))

                # 坐标映射：csv(row,col) → Word "table{ti+1} row{r+2} col{c+1}"
                coord_map = {}
                for ri in range(len(data_rows)):
                    for ci in range(len(headers)):
                        coord_map[f"{ri},{ci}"] = f"table{ti+1}_row{ri+2}_col{ci+1}"

                tables.append(TableBlock(
                    meta=TableMeta(
                        table_id=table_id,
                        source_sheet="",
                        source_range=f"table{ti+1}",
                        row_count=len(data_rows),
                        col_count=len(headers),
                        coord_map=coord_map,
                        merged_cells=merged,
                    ),
                    headers=headers,
                    rows=data_rows,
                ))
            except Exception as exc:
                warnings.append(Warning(
                    code="TABLE_PARSE_ERROR",
                    message=f"Table {ti+1} parsing error: {exc}",
                    loc=Loc(type=LocType.PARAGRAPH.value, value=0,
                            extra={"table_index": ti}),
                ))

        return tables

    @staticmethod
    def _detect_merged_cells(table) -> list[str]:
        """检测 python-docx 表格中的合并区域"""
        merged = set()
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # python-docx 中合并的单元格 span 信息可通过 _tc 获取
                tc = cell._tc
                grid_span = tc.get(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan'
                )
                v_merge = tc.find(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge'
                )
                if grid_span is not None:
                    span = int(grid_span.get(
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '1'
                    ))
                    if span > 1:
                        merged.add(f"row{row_idx+1}_col{col_idx+1}_colspan{span}")
                if v_merge is not None:
                    merged.add(f"row{row_idx+1}_col{col_idx+1}_vmerge")
        return sorted(merged)

    # ------------------------------------------------------------------
    # 图片
    # ------------------------------------------------------------------
    def _extract_images(
        self, doc, fpath: Path, warnings: list[Warning]
    ) -> list[ImageBlock]:
        images: list[ImageBlock] = []
        try:
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.reltype:
                    images.append(ImageBlock(
                        image_id=rel_id,
                        caption="",
                        path=rel.target_ref,
                        loc=Loc(type=LocType.PARAGRAPH.value, value=0,
                                extra={"rel_id": rel_id}),
                    ))
        except Exception as exc:
            warnings.append(Warning(
                code="IMAGE_EXTRACT_FAILED",
                message=f"Image extraction failed: {exc}",
                loc=Loc(type=LocType.PARAGRAPH.value, value=0),
            ))
        return images

    # ------------------------------------------------------------------
    # 辅助
    # ------------------------------------------------------------------
    @staticmethod
    def _detect_title(doc) -> str:
        for para in doc.paragraphs:
            style = para.style.name.lower() if para.style else ""
            text = para.text.strip()
            if text and ("title" in style or "标题" in style):
                return text
        # fallback: 第一个非空行
        for para in doc.paragraphs:
            if para.text.strip():
                return para.text.strip()
        return ""

    def _source_dict(self, fpath: Path) -> dict:
        return {
            "file_name": fpath.name,
            "file_type": "docx",
            "file_size_bytes": fpath.stat().st_size,
            "file_path": str(fpath),
        }